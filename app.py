import streamlit as st
import pandas as pd
import io

# Impor modul secara aman agar aplikasi tidak crash
try:
    from st_gsheets_connection import GSheetsConnection
    GSHEETS_OK = True
except ImportError:
    GSHEETS_OK = False

try:
    import google.generativeai as genai
    GEMINI_OK = True
except ImportError:
    GEMINI_OK = False

try:
    from docx import Document
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# Konfigurasi Tampilan
st.set_page_config(page_title="Projek ADM", page_icon="📝", layout="wide")
st.title("📝 Sistem Administrasi & Pembuat Surat Otomatis")
st.write("---")

# 1. Fitur Google Sheets
st.subheader("📊 Data Google Sheets")
if GSHEETS_OK:
    try:
        conn = st.connection("gsheets", type=GSheetsConnection)
        df = conn.read(ttl=0)
        st.success("✅ Terhubung ke Google Sheets!")
        with st.expander("Lihat Data"):
            st.dataframe(df)
    except Exception as e:
        st.warning(f"⚠️ Gagal membaca Google Sheets: {e}")
else:
    st.error("❌ Modul `st-gsheets-connection` belum terpasang di server.")

# 2. Fitur Gemini AI
st.write("---")
st.subheader("🤖 Generator Surat (Gemini AI)")
if GEMINI_OK:
    gemini_key = st.secrets.get("GEMINI_API_KEY", None)
    if gemini_key:
        try:
            genai.configure(api_key=gemini_key)
            model = genai.GenerativeModel("gemini-1.5-flash")
            prompt_user = st.text_area("Instruksi Surat:", "Buatkan draf surat permohonan izin...")
            
            if st.button("Generate Teks Surat"):
                with st.spinner("Sedang memproses..."):
                    res = model.generate_content(prompt_user)
                    st.session_state["draft_surat"] = res.text
                    st.success("✅ Draf berhasil dibuat!")
                    st.write(res.text)
        except Exception as e:
            st.error(f"❌ Kendala Gemini AI: {e}")
    else:
        st.warning("⚠️ GEMINI_API_KEY belum ditambahkan di Streamlit Secrets.")
else:
    st.error("❌ Modul `google-generativeai` belum terpasang di server.")

# 3. Fitur Export Word
st.write("---")
st.subheader("📄 Unduh Dokumen Word (.docx)")
if DOCX_OK:
    if "draft_surat" in st.session_state and st.session_state["draft_surat"]:
        if st.button("Buat File Word (.docx)"):
            doc = Document()
            doc.add_heading("DRAF SURAT RESMI", level=1)
            doc.add_paragraph(st.session_state["draft_surat"])
            
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            
            st.download_button(
                label="📥 Download Surat (.docx)",
                data=buf,
                file_name="Surat_Otomatis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.info("Buat draf surat menggunakan Gemini AI terlebih dahulu.")
else:
    st.error("❌ Modul `python-docx` belum terpasang di server.")